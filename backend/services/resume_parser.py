"""
Resume Parsing Service
Extracts text, skills, and experience level from resume files.
Uses Gemini AI for smart extraction when available, falls back to rule-based methods.
"""

import json
import re
import os
from typing import List, Optional, Any

import config

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """Normalize whitespace, strip null bytes and control chars."""
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(file_path: str, filename: str) -> str:
    """Extract text from PDF, DOCX, or TXT file with improved handling."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    text = ""

    if ext == "pdf":
        text = _extract_pdf(file_path)
    elif ext in ("docx", "doc"):
        text = _extract_docx(file_path)
    elif ext == "txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            print(f"TXT parsing error: {e}")

    return _clean_text(text)


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2 with per-page error handling."""
    text = ""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    print(f"  PDF page {i+1}/{total_pages}: no text extracted")
            except Exception as e:
                print(f"  PDF page {i+1}/{total_pages} error: {e}")

        char_count = len(text.strip())
        if total_pages > 0 and char_count < 50 * total_pages:
            print(f"  Warning: PDF has {total_pages} pages but only {char_count} chars extracted")

    except Exception as e:
        print(f"PDF parsing error: {e}")

    return text


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX including paragraphs and tables."""
    text = ""
    try:
        from docx import Document
        doc = Document(file_path)

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Also extract text from tables (common in formatted resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"

    except Exception as e:
        print(f"DOCX parsing error: {e}")

    return text


# ---------------------------------------------------------------------------
# Skill matching
# ---------------------------------------------------------------------------

SKILL_ALIASES = {
    "react": ["react", "reactjs", "react.js"],
    "node.js": ["node.js", "nodejs", "node js"],
    "vue": ["vue", "vuejs", "vue.js"],
    "angular": ["angular", "angularjs", "angular.js"],
    "typescript": ["typescript", "ts"],
    "javascript": ["javascript", "js", "ecmascript", "es6"],
    "python": ["python", "python3"],
    "c++": ["c++", "cpp", "cplusplus"],
    "c#": ["c#", "csharp", "c sharp"],
    ".net": [".net", "dotnet", "dot net"],
    "postgresql": ["postgresql", "postgres", "psql"],
    "mongodb": ["mongodb", "mongo"],
    "mysql": ["mysql"],
    "aws": ["aws", "amazon web services"],
    "gcp": ["gcp", "google cloud", "google cloud platform"],
    "azure": ["azure", "microsoft azure"],
    "docker": ["docker"],
    "kubernetes": ["kubernetes", "k8s"],
    "ci/cd": ["ci/cd", "cicd", "ci cd", "continuous integration"],
    "rest": ["rest", "restful", "rest api", "restful api"],
    "graphql": ["graphql", "graph ql"],
    "machine learning": ["machine learning", "ml"],
    "sql": ["sql"],
    "java": ["java"],
    "go": ["golang", "go lang"],
    "ruby": ["ruby", "ruby on rails", "rails"],
    "php": ["php"],
    "swift": ["swift"],
    "kotlin": ["kotlin"],
    "rust": ["rust"],
    "next.js": ["next.js", "nextjs", "next js"],
    "express": ["express", "expressjs", "express.js"],
    "spring": ["spring", "spring boot", "springboot"],
    "django": ["django"],
    "flask": ["flask"],
    "fastapi": ["fastapi", "fast api"],
    "redux": ["redux"],
    "html": ["html", "html5"],
    "css": ["css", "css3"],
    "sass": ["sass", "scss"],
    "git": ["git", "github", "gitlab"],
    "linux": ["linux", "ubuntu", "centos"],
    "terraform": ["terraform"],
    "figma": ["figma"],
    "jira": ["jira"],
}


def _get_skill_variants(skill: str) -> List[str]:
    """Get all variant forms of a skill name."""
    skill_lower = skill.lower().strip()
    # Check alias map
    for canonical, aliases in SKILL_ALIASES.items():
        if skill_lower in aliases or skill_lower == canonical:
            return aliases
    # No alias found - return the skill itself
    return [skill_lower]


def _skill_in_text(skill_variants: List[str], text_lower: str) -> bool:
    """Check if any skill variant appears in text using word boundaries."""
    for variant in skill_variants:
        # For very short terms (1-2 chars like "C", "R", "Go"), require strict word boundaries
        escaped = re.escape(variant)
        pattern = r'\b' + escaped + r'\b'
        try:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        except re.error:
            # Fallback for tricky patterns
            if variant in text_lower:
                return True
    return False


def extract_skills(resume_text: str, job_skills: List[str]) -> List[str]:
    """Extract matching skills using normalized regex matching."""
    if not resume_text or not job_skills:
        return []

    text_lower = resume_text.lower()
    matched = []
    for skill in job_skills:
        variants = _get_skill_variants(skill)
        if _skill_in_text(variants, text_lower):
            matched.append(skill)

    return matched


# ---------------------------------------------------------------------------
# Experience level
# ---------------------------------------------------------------------------

def determine_experience_level(resume_text: str, experience_years: Optional[int] = None):
    """Determine candidate experience level from resume content and stated years."""
    from models import ExperienceLevel

    # Tier 1: Rule-based
    years = experience_years

    # If years not provided, try to extract from resume text
    if years is None and resume_text:
        years = _extract_years_from_text(resume_text)

    if years is not None:
        if years <= 2:
            return ExperienceLevel.JUNIOR
        elif years <= 5:
            return ExperienceLevel.MID
        else:
            return ExperienceLevel.SENIOR

    # Check for title keywords
    if resume_text:
        text_lower = resume_text.lower()
        senior_keywords = ["senior", "lead", "principal", "architect", "director", "manager", "head of", "vp ", "vice president"]
        junior_keywords = ["intern", "trainee", "entry level", "entry-level", "junior", "fresher", "graduate"]

        for kw in senior_keywords:
            if kw in text_lower:
                return ExperienceLevel.SENIOR
        for kw in junior_keywords:
            if kw in text_lower:
                return ExperienceLevel.JUNIOR

        # Default to MID if we have resume text but can't determine
        return ExperienceLevel.MID

    return None


def _extract_years_from_text(text: str) -> Optional[int]:
    """Extract years of experience from resume text using regex."""
    patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)\s*(?:of\s+)?(\d+)\+?\s*(?:years?|yrs?)',
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:in\s+)?(?:software|development|engineering|programming|IT)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                return int(match.group(1))
            except (ValueError, IndexError):
                pass
    return None


# ---------------------------------------------------------------------------
# Contact info extraction (name, email, phone, current position)
# ---------------------------------------------------------------------------

def _extract_email(text: str) -> Optional[str]:
    """Extract email address from resume text."""
    match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return match.group(0) if match else None


def _extract_phone(text: str) -> Optional[str]:
    """Extract phone number from resume text."""
    patterns = [
        r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'(?:\+?\d{1,3}[-.\s]?)?\d{5}[-.\s]?\d{5}',
        r'(?:\+91[-.\s]?)?\d{10}',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def _extract_name(text: str) -> Optional[str]:
    """Extract candidate name from resume - typically the first non-empty line."""
    lines = text.strip().split('\n')
    for line in lines[:5]:
        line = line.strip()
        if not line or len(line) < 2:
            continue
        # Skip lines that look like emails, phones, URLs, or common headers
        if re.search(r'@|http|www\.|\.com|\.in|\.org|resume|curriculum|vitae|cv\b|\d{5,}|education|experience|summary|objective|skills', line, re.IGNORECASE):
            continue
        # Clean pipe-separated headers like "Deep M. Mehta | Full Stack Developer site.com"
        # Take only the name part (before |)
        if '|' in line:
            line = line.split('|')[0].strip()
        # A name is typically 2-4 words, all alpha (with dots for initials)
        words = line.split()
        if 1 <= len(words) <= 5 and all(re.match(r'^[a-zA-Z.\'-]+$', w) for w in words):
            return line
    return None


def _extract_current_position(text: str) -> Optional[str]:
    """Extract current job title/position from resume text."""
    text_lower = text.lower()
    # Look for explicit current position patterns
    patterns = [
        r'(?:current(?:ly)?|present)\s*(?:role|position|title|designation)\s*[:\-–]\s*(.+)',
        r'(?:working as|currently working as|role)\s*[:\-–]?\s*(.+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text_lower)
        if match:
            pos = match.group(1).strip().rstrip('.').strip()
            if 3 < len(pos) < 80:
                start = match.start(1)
                return text[start:start + len(pos)].strip()

    # Look for common title keywords in the header area
    title_keywords = [
        "software engineer", "developer", "frontend developer", "backend developer",
        "full stack developer", "fullstack developer", "data scientist", "data analyst",
        "data engineer", "devops engineer", "ui designer", "ux designer", "ui/ux designer",
        "product manager", "project manager", "qa engineer", "tester", "analyst",
        "consultant", "architect", "team lead", "tech lead", "manager", "intern", "trainee",
    ]
    # Check first ~500 chars for title keywords
    top_text = text_lower[:500]
    for kw in title_keywords:
        if kw in top_text:
            # Extract just the title, not the whole line (which may have names/URLs)
            # If the line has a pipe separator, look for the part with the keyword
            for line in text[:500].split('\n'):
                if kw in line.lower():
                    # If pipe-separated, take only the part with the keyword
                    if '|' in line:
                        for part in line.split('|'):
                            if kw in part.lower():
                                clean = part.strip()
                                # Remove URLs and emails
                                clean = re.sub(r'https?://\S+|www\.\S+|\S+\.\w{2,3}\.\w{2,3}|\S+@\S+', '', clean).strip()
                                if 3 < len(clean) < 60:
                                    return clean
                    else:
                        clean = line.strip()
                        clean = re.sub(r'https?://\S+|www\.\S+|\S+\.\w{2,3}\.\w{2,3}|\S+@\S+', '', clean).strip()
                        if 3 < len(clean) < 60:
                            return clean
    return None


def extract_contact_info(resume_text: str) -> dict:
    """Extract contact information from resume text."""
    return {
        "name": _extract_name(resume_text),
        "email": _extract_email(resume_text),
        "phone": _extract_phone(resume_text),
        "current_position": _extract_current_position(resume_text),
        "experience_years": _extract_years_from_text(resume_text),
    }


# ---------------------------------------------------------------------------
# Gemini AI integration (single combined call for skills + experience)
# ---------------------------------------------------------------------------

def _analyze_with_gemini(resume_text: str, job_skills: List[str]) -> Optional[dict]:
    """Use Gemini AI to extract skills and experience level from resume."""
    try:
        from google import genai
        client = genai.Client(api_key=config.GEMINI_API_KEY)

        # Truncate resume text to stay within token limits
        truncated = resume_text[:3000]
        skills_list = ", ".join(job_skills)

        prompt = f"""You are a resume analysis expert. Given the following resume text and required job skills, perform these tasks:

1. SKILL MATCHING: Identify which of the required skills the candidate possesses. Consider synonyms, abbreviations, and related technologies. For example, "ReactJS" matches "React", "Node" matches "Node.js", etc. Only return skills from the provided list.

2. ALL SKILLS: Extract ALL technical skills, tools, frameworks, and technologies mentioned in the resume (not just from the required list). Include programming languages, frameworks, databases, tools, cloud platforms, etc.

3. EXPERIENCE LEVEL: Determine the candidate's experience level as one of: "Junior" (0-2 years), "Mid" (3-5 years), or "Senior" (6+ years).

4. CONTACT INFO: Extract the candidate's full name, email address, phone number, current job title/position, and total years of experience (as integer).

RESUME TEXT:
{truncated}

REQUIRED SKILLS:
{skills_list}

Respond ONLY with valid JSON (no markdown, no explanation):
{{"matched_skills": ["skill1", "skill2"], "all_skills": ["React", "Python", "Docker", "..."], "experience_level": "Junior", "name": "John Doe", "email": "john@example.com", "phone": "+1234567890", "current_position": "Software Engineer", "experience_years": 3}}"""

        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )

        # Parse response
        result = _extract_json_from_response(response.text)
        if result and isinstance(result, dict):
            return result

    except Exception as e:
        print(f"Gemini resume analysis failed (falling back to rule-based): {e}")

    return None


def _extract_json_from_response(text: str) -> Any:
    """Extract JSON from LLM response text."""
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    match = re.search(r'```(?:json)?\s*\n?([\s\S]*?)\n?```', text)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    for pattern in [r'(\{[\s\S]*\})']:
        match = re.search(pattern, text)
        if match:
            try:
                return json.loads(match.group(1))
            except json.JSONDecodeError:
                continue
    return None


# ---------------------------------------------------------------------------
# Main orchestrator
# ---------------------------------------------------------------------------

def parse_resume(
    file_path: str,
    filename: str,
    job_skills: List[str],
    experience_years: Optional[int] = None
) -> dict:
    """
    Full resume parsing pipeline.
    Returns dict with: parsed_text, skills, experience_level, parsing_status
    """
    # Step 1: Extract text
    parsed_text = extract_text(file_path, filename)

    # Determine parsing status
    if not parsed_text:
        parsing_status = "failed"
    elif len(parsed_text) < 100:
        parsing_status = "partial"
    else:
        parsing_status = "completed"

    # Step 2: Extract skills (matched against job requirements)
    skills = extract_skills(parsed_text, job_skills)

    # Also extract general skills from resume (using all known skill aliases)
    if parsed_text:
        all_known_skills = list(SKILL_ALIASES.keys())
        general_skills = extract_skills(parsed_text, all_known_skills)
        for s in general_skills:
            if s not in skills:
                skills.append(s)

    # Step 3: Determine experience level
    exp_level = determine_experience_level(parsed_text, experience_years)

    # Step 4: Extract contact info (rule-based)
    contact_info = extract_contact_info(parsed_text) if parsed_text else {}

    # Step 5: Gemini AI enhancement
    ai_result = None
    if config.GEMINI_API_KEY and parsed_text and job_skills:
        ai_result = _analyze_with_gemini(parsed_text, job_skills)
        if ai_result:
            # Merge AI skills
            if "matched_skills" in ai_result:
                for s in ai_result["matched_skills"]:
                    s_lower = s.lower()
                    for job_skill in job_skills:
                        if job_skill.lower() == s_lower and job_skill not in skills:
                            skills.append(job_skill)
                            break
            # Use AI experience level if rule-based didn't find one
            if exp_level is None and "experience_level" in ai_result:
                from models import ExperienceLevel
                level_map = {"junior": ExperienceLevel.JUNIOR, "mid": ExperienceLevel.MID, "senior": ExperienceLevel.SENIOR}
                exp_level = level_map.get(ai_result["experience_level"].lower())

            # Use AI contact info as fallback for rule-based extraction
            for field in ("name", "email", "phone", "current_position", "experience_years"):
                if not contact_info.get(field) and ai_result.get(field):
                    contact_info[field] = ai_result[field]

            # Merge all_skills from AI (general skills found in resume)
            if "all_skills" in ai_result and isinstance(ai_result["all_skills"], list):
                for s in ai_result["all_skills"]:
                    if s and s not in skills:
                        skills.append(s)

    return {
        "parsed_text": parsed_text,
        "skills": skills,
        "experience_level": exp_level,
        "parsing_status": parsing_status,
        "contact_info": contact_info,
    }
